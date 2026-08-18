from docx import Document
from docx.shared import Pt, Inches
import base64
import csv
import io

def _test_cases_by_fr(prd):
    """Groups test_cases by the FR they verify, in functional_requirements
    order -- with any test case whose fr_id doesn't match a known FR (should
    not normally happen) appended at the end so nothing silently drops."""
    by_fr: dict[str, list] = {}
    for tc in prd.test_cases:
        by_fr.setdefault(tc.fr_id, []).append(tc)
    ordered: list[tuple[str, list]] = []
    for fr in prd.functional_requirements:
        group = by_fr.pop(fr.id, None)
        if group:
            ordered.append((fr.id, group))
    ordered.extend(by_fr.items())
    return ordered


def export_to_markdown(prd) -> str:
    lines = [f"# {prd.title}\n"]

    has_overview = (
        prd.problem_statement or prd.goals or prd.target_users
        or prd.personas or prd.use_cases or prd.open_questions
    )
    if has_overview:
        lines.append("## Overview")
        if prd.problem_statement:
            lines.append(f"{prd.problem_statement}\n")
        if prd.goals:
            lines.append("**Goals**:")
            for g in prd.goals:
                lines.append(f"- {g}")
            lines.append("")
        if prd.target_users:
            lines.append("**Target users**:")
            for u in prd.target_users:
                lines.append(f"- {u}")
            lines.append("")
        if prd.personas:
            lines.append("### Personas")
            for p in prd.personas:
                lines.append(f"- **{p.name}**: {p.description}")
                for pp in p.pain_points:
                    lines.append(f"  - Pain point: {pp}")
            lines.append("")
        if prd.use_cases:
            lines.append("### Use Cases")
            for u in prd.use_cases:
                lines.append(f"- **{u.title}** ({u.actor}): {u.goal} -- triggered by {u.trigger}")
            lines.append("")
        if prd.open_questions:
            lines.append("### Open Questions")
            for q in prd.open_questions:
                lines.append(f"- {q}")
            lines.append("")

    lines.append("## User Scenarios & Testing")
    for story in prd.user_stories:
        lines.append(f"### {story.title} (Priority: {story.priority})")
        lines.append(f"{story.description}\n")
        lines.append(f"**Why this priority**: {story.why_this_priority}\n")
        lines.append(f"**Independent Test**: {story.independent_test}\n")
        if story.acceptance_scenarios:
            lines.append("**Acceptance Scenarios**:")
            for i, sc in enumerate(story.acceptance_scenarios, 1):
                lines.append(
                    f"{i}. **Given** {sc.given}, **When** {sc.when}, "
                    f"**Then** {sc.then}"
                )
        lines.append("")

    lines.append("### Edge Cases")
    for edge in prd.edge_cases:
        lines.append(f"- {edge}")
    lines.append("")

    lines.append("## Requirements")
    lines.append("### Functional Requirements")
    for fr in prd.functional_requirements:
        lines.append(f"- **{fr.id}** [{fr.kind}]: {fr.text}")
    lines.append("")

    if prd.key_entities:
        lines.append("### Key Entities")
        for ke in prd.key_entities:
            lines.append(f"- **{ke.name}**: {ke.description}")
        lines.append("")

    lines.append("## Test Cases")
    for fr_id, group in _test_cases_by_fr(prd):
        lines.append(f"### {fr_id}")
        for tc in group:
            lines.append(f"- **{tc.id}** {tc.title}: Given {tc.given}, When {tc.when}, Then {tc.then}")
    lines.append("")

    lines.append("## Assumptions")
    for a in prd.assumptions:
        lines.append(f"- {a}")

    if prd.architecture_decisions or prd.technical_context or prd.architecture_clarify_history:
        lines.append("")
        lines.append("## Architecture Decisions")
        if prd.technical_context:
            lines.append("**Technical context**:")
            for h in prd.technical_context:
                lines.append(f"- {h.question.question} -> {h.answer}")
            lines.append("")
        if prd.architecture_clarify_history:
            lines.append("**Architecture clarifications**:")
            for h in prd.architecture_clarify_history:
                lines.append(f"- {h.question.question} -> {h.answer}")
            lines.append("")
        for adr in prd.architecture_decisions:
            lines.append(f"### {adr.id}: {adr.title}")
            lines.append(f"**Status**: {adr.status}\n")
            lines.append(f"**Context**: {adr.context}\n")
            lines.append(f"**Decision**: {adr.decision}\n")
            lines.append(f"**Consequences**: {adr.consequences}\n")

    if prd.epics:
        lines.append("")
        lines.append("## Epics")
        for epic in prd.epics:
            jira_note = f" ({epic.jira_key})" if epic.jira_key else ""
            lines.append(f"### {epic.id}: {epic.title}{jira_note} [{epic.business_impact} impact]")
            lines.append(f"{epic.description}\n")
            for story in epic.stories:
                story_jira = f" ({story.jira_key})" if story.jira_key else ""
                lines.append(f"- **[{story.story_type}] {story.id}**: {story.title}{story_jira}")
                lines.append(f"  {story.description}")
                for ac in story.acceptance_criteria:
                    lines.append(f"  - {ac}")
                if story.notes:
                    lines.append(f"  *Note: {story.notes}*")
            lines.append("")

    if prd.diagrams:
        lines.append("")
        lines.append("## Diagrams")
        for diagram in prd.diagrams:
            lines.append(f"### {diagram.title}")
            lines.append("```mermaid")
            lines.append(diagram.mermaid_source)
            lines.append("```")
            lines.append("")

    if prd.completeness_review and prd.completeness_review.issues:
        lines.append("")
        lines.append("## Completeness Review")
        for issue in prd.completeness_review.issues:
            related = f" ({', '.join(issue.related_ids)})" if issue.related_ids else ""
            lines.append(f"- **[{issue.severity}] {issue.area}**: {issue.description}{related}")
        lines.append("")

    return "\n".join(lines)


def export_to_docx(prd, path: str):
    doc = Document()
    doc.add_heading(prd.title, level=0)

    has_overview = (
        prd.problem_statement or prd.goals or prd.target_users
        or prd.personas or prd.use_cases or prd.open_questions
    )
    if has_overview:
        doc.add_heading("Overview", level=1)
        if prd.problem_statement:
            doc.add_paragraph(prd.problem_statement)
        if prd.goals:
            p = doc.add_paragraph()
            p.add_run("Goals").bold = True
            for g in prd.goals:
                doc.add_paragraph(g, style="List Bullet")
        if prd.target_users:
            p = doc.add_paragraph()
            p.add_run("Target users").bold = True
            for u in prd.target_users:
                doc.add_paragraph(u, style="List Bullet")
        if prd.personas:
            doc.add_heading("Personas", level=2)
            for persona in prd.personas:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{persona.name}: ").bold = True
                p.add_run(persona.description)
                for pp in persona.pain_points:
                    doc.add_paragraph(f"Pain point: {pp}", style="List Bullet 2")
        if prd.use_cases:
            doc.add_heading("Use Cases", level=2)
            for uc in prd.use_cases:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{uc.title} ({uc.actor}): ").bold = True
                p.add_run(f"{uc.goal} -- triggered by {uc.trigger}")
        if prd.open_questions:
            doc.add_heading("Open Questions", level=2)
            for q in prd.open_questions:
                doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("User Scenarios & Testing", level=1)
    for story in prd.user_stories:
        doc.add_heading(f"{story.title} (Priority: {story.priority})", level=2)
        doc.add_paragraph(story.description)

        p = doc.add_paragraph()
        p.add_run("Why this priority: ").bold = True
        p.add_run(story.why_this_priority)

        p = doc.add_paragraph()
        p.add_run("Independent Test: ").bold = True
        p.add_run(story.independent_test)

        if story.acceptance_scenarios:
            doc.add_paragraph("Acceptance Scenarios:", style="Intense Quote")
            for sc in story.acceptance_scenarios:
                doc.add_paragraph(
                    f"Given {sc.given}, When {sc.when}, Then {sc.then}",
                    style="List Number",
                )

    doc.add_heading("Edge Cases", level=2)
    for edge in prd.edge_cases:
        doc.add_paragraph(edge, style="List Bullet")

    doc.add_heading("Requirements", level=1)
    doc.add_heading("Functional Requirements", level=2)
    for fr in prd.functional_requirements:
        doc.add_paragraph(f"{fr.id} [{fr.kind}]: {fr.text}", style="List Bullet")

    if prd.key_entities:
        doc.add_heading("Key Entities", level=2)
        for ke in prd.key_entities:
            doc.add_paragraph(f"{ke.name}: {ke.description}", style="List Bullet")

    doc.add_heading("Test Cases", level=1)
    for fr_id, group in _test_cases_by_fr(prd):
        doc.add_heading(fr_id, level=2)
        for tc in group:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{tc.id} {tc.title}: ").bold = True
            p.add_run(f"Given {tc.given}, When {tc.when}, Then {tc.then}")

    doc.add_heading("Assumptions", level=1)
    for a in prd.assumptions:
        doc.add_paragraph(a, style="List Bullet")

    if prd.architecture_decisions or prd.technical_context or prd.architecture_clarify_history:
        doc.add_heading("Architecture Decisions", level=1)
        if prd.technical_context:
            p = doc.add_paragraph()
            p.add_run("Technical context").bold = True
            for h in prd.technical_context:
                doc.add_paragraph(f"{h.question.question} -> {h.answer}", style="List Bullet")
        if prd.architecture_clarify_history:
            p = doc.add_paragraph()
            p.add_run("Architecture clarifications").bold = True
            for h in prd.architecture_clarify_history:
                doc.add_paragraph(f"{h.question.question} -> {h.answer}", style="List Bullet")
        for adr in prd.architecture_decisions:
            doc.add_heading(f"{adr.id}: {adr.title}", level=2)
            p = doc.add_paragraph()
            p.add_run("Status: ").bold = True
            p.add_run(adr.status)
            p = doc.add_paragraph()
            p.add_run("Context: ").bold = True
            p.add_run(adr.context)
            p = doc.add_paragraph()
            p.add_run("Decision: ").bold = True
            p.add_run(adr.decision)
            p = doc.add_paragraph()
            p.add_run("Consequences: ").bold = True
            p.add_run(adr.consequences)

    if prd.epics:
        doc.add_heading("Epics", level=1)
        for epic in prd.epics:
            title = f"{epic.id}: {epic.title}"
            if epic.jira_key:
                title += f" ({epic.jira_key})"
            title += f" [{epic.business_impact} impact]"
            doc.add_heading(title, level=2)
            doc.add_paragraph(epic.description)
            for story in epic.stories:
                summary = f"[{story.story_type}] {story.id}: {story.title}"
                if story.jira_key:
                    summary += f" ({story.jira_key})"
                doc.add_paragraph(summary, style="List Bullet")
                if story.description:
                    doc.add_paragraph(story.description, style="List Bullet 2")
                for ac in story.acceptance_criteria:
                    doc.add_paragraph(ac, style="List Bullet 2")
                if story.notes:
                    note_p = doc.add_paragraph(style="List Bullet 2")
                    note_p.add_run(f"Note: {story.notes}").italic = True

    if prd.diagrams:
        doc.add_heading("Diagrams", level=1)
        if any(not d.png_base64 for d in prd.diagrams):
            note = doc.add_paragraph()
            note.add_run(
                "Diagrams below without an embedded image are Mermaid source -- "
                "paste into a Mermaid-aware renderer (or the Mermaid Chart "
                "Office add-in) to view; Word does not render Mermaid natively."
            ).italic = True
        for diagram in prd.diagrams:
            doc.add_heading(diagram.title, level=2)
            if diagram.png_base64:
                image_bytes = base64.b64decode(diagram.png_base64)
                doc.add_picture(io.BytesIO(image_bytes), width=Inches(6))
            else:
                p = doc.add_paragraph()
                source_lines = diagram.mermaid_source.splitlines()
                for i, line in enumerate(source_lines):
                    run = p.add_run(line)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    if i < len(source_lines) - 1:
                        run.add_break()

    if prd.completeness_review and prd.completeness_review.issues:
        doc.add_heading("Completeness Review", level=1)
        for issue in prd.completeness_review.issues:
            related = f" ({', '.join(issue.related_ids)})" if issue.related_ids else ""
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{issue.severity}] {issue.area}: ").bold = True
            p.add_run(f"{issue.description}{related}")

    doc.save(path)


def export_stories_to_csv(prd, path: str):
    """JIRA-importable CSV bridge (no API integration needed yet)."""
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Priority", "Summary", "Description", "Acceptance Criteria"])
        for story in prd.user_stories:
            criteria = "; ".join(
                f"Given {sc.given}, When {sc.when}, Then {sc.then}"
                for sc in story.acceptance_scenarios
            )
            writer.writerow([story.priority, story.title, story.description, criteria])


def export_brief_to_markdown(brief) -> str:
    """Standalone audience-brief deliverable -- a projection of the spec, kept
    separate from the spec's own export."""
    lines = [f"# {brief.title}\n"]
    for section in brief.sections:
        lines.append(f"## {section.heading}")
        lines.append(f"{section.body}\n")
    if brief.key_asks:
        lines.append("## Asks")
        for ask in brief.key_asks:
            lines.append(f"- {ask}")
    return "\n".join(lines)


def export_brief_to_docx(brief, path: str):
    doc = Document()
    doc.add_heading(brief.title, level=0)
    for section in brief.sections:
        doc.add_heading(section.heading, level=1)
        for para in section.body.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("- "):
                doc.add_paragraph(para[2:], style="List Bullet")
            else:
                doc.add_paragraph(para)
    if brief.key_asks:
        doc.add_heading("Asks", level=1)
        for ask in brief.key_asks:
            doc.add_paragraph(ask, style="List Bullet")
    doc.save(path)


def export_update_to_markdown(update) -> str:
    """Standalone stakeholder-update deliverable composed from version diffs."""
    lines = [f"# {update.title}\n"]
    lines.append(f"*{update.summary}*\n")
    for section in update.sections:
        lines.append(f"## {section.heading}")
        lines.append(f"{section.body}\n")
    if update.decisions_needed:
        lines.append("## Decisions needed")
        for d in update.decisions_needed:
            lines.append(f"- {d}")
    return "\n".join(lines)


def export_update_to_docx(update, path: str):
    doc = Document()
    doc.add_heading(update.title, level=0)
    summary = doc.add_paragraph()
    summary.add_run(update.summary).italic = True
    for section in update.sections:
        doc.add_heading(section.heading, level=1)
        for para in section.body.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("- "):
                doc.add_paragraph(para[2:], style="List Bullet")
            else:
                doc.add_paragraph(para)
    if update.decisions_needed:
        doc.add_heading("Decisions needed", level=1)
        for d in update.decisions_needed:
            doc.add_paragraph(d, style="List Bullet")
    doc.save(path)


def export_epics_to_csv(prd, path: str):
    """Jira bulk-CSV-importer bridge for Epics + Stories (offline alternative
    to the live push in jira_client.py -- Jira's importer links a Story to
    its Epic by matching 'Epic Link' text against an Epic row's 'Epic Name')."""
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Issue Type", "Epic Name", "Epic Link", "Summary", "Description", "Acceptance Criteria"])
        for epic in prd.epics:
            writer.writerow(["Epic", epic.title, "", epic.title, epic.description, ""])
            for story in epic.stories:
                criteria = "; ".join(story.acceptance_criteria)
                writer.writerow(
                    ["Story", "", epic.title, story.title, story.description, criteria]
                )
